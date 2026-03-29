"""
Document Parser for the Offline Policy Gap Analyzer.

This module provides the DocumentParser class for extracting text from policy
documents (PDF, DOCX, TXT) while preserving structural boundaries.
"""

import logging
from pathlib import Path
from typing import Optional
import chardet

# PDF parsing
import fitz  # PyMuPDF
import pdfplumber

# DOCX parsing
from docx import Document as DocxDocument

from models.domain import (
    ParsedDocument,
    DocumentStructure,
    Heading,
    Section,
    Paragraph,
)


# Custom exceptions
class UnsupportedFormatError(Exception):
    """Raised when file format is not supported."""
    pass


class OCRRequiredError(Exception):
    """Raised when PDF contains only scanned images without text layer."""
    pass


class ParsingError(Exception):
    """Raised when document parsing fails."""
    pass


logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Parser for extracting text from policy documents.
    
    Supports PDF (via PyMuPDF with pdfplumber fallback), DOCX, and TXT formats.
    Preserves document structure including headings, sections, and paragraphs.
    """
    
    SUPPORTED_FORMATS = ['pdf', 'docx', 'txt', 'md']
    MAX_PAGES = 100
    
    def __init__(self):
        """Initialize the DocumentParser."""
        self.logger = logging.getLogger(__name__)
    
    def parse(self, file_path: str, file_type: Optional[str] = None) -> ParsedDocument:
        """
        Parse document and extract structured text.
        
        Args:
            file_path: Path to policy document
            file_type: One of 'pdf', 'docx', 'txt' (auto-detected if None)
            
        Returns:
            ParsedDocument with text, metadata, and structure
            
        Raises:
            UnsupportedFormatError: If file type not supported
            OCRRequiredError: If PDF contains only scanned images
            ParsingError: If extraction fails
        """
        path = Path(file_path)
        
        if not path.exists():
            raise ParsingError(f"File not found: {file_path}")
        
        # Auto-detect file type if not provided
        if file_type is None:
            file_type = path.suffix.lstrip('.').lower()
        
        file_type = file_type.lower()
        
        if file_type not in self.SUPPORTED_FORMATS:
            raise UnsupportedFormatError(
                f"Unsupported format '{file_type}'. "
                f"Supported formats: {', '.join(self.SUPPORTED_FORMATS)}"
            )
        
        # Route to appropriate parser
        if file_type == 'pdf':
            return self._parse_pdf(file_path)
        elif file_type == 'docx':
            return self._parse_docx(file_path)
        elif file_type in ('txt', 'md'):
            return self._parse_txt(file_path)
    
    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """
        Parse PDF document using PyMuPDF with pdfplumber fallback.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            ParsedDocument with extracted text and structure
            
        Raises:
            OCRRequiredError: If PDF is scanned without text layer
            ParsingError: If extraction fails
        """
        try:
            # Try PyMuPDF first (primary parser)
            doc = fitz.open(file_path)
            page_count = len(doc)
            
            # Enforce page limit
            if page_count > self.MAX_PAGES:
                raise ParsingError(
                    f"Document exceeds maximum page limit of {self.MAX_PAGES} pages "
                    f"(found {page_count} pages)"
                )
            
            # Extract text from all pages
            text_parts = []
            has_text = False
            
            for page_num in range(page_count):
                page = doc[page_num]
                page_text = page.get_text()
                
                if page_text.strip():
                    has_text = True
                    text_parts.append(page_text)
            
            doc.close()
            
            # Check if PDF has text layer (not scanned)
            if not has_text:
                raise OCRRequiredError(
                    "OCR is not supported; please provide a text-based PDF"
                )
            
            full_text = '\n'.join(text_parts)
            
            # Try to extract structure
            structure = self._extract_pdf_structure(file_path, full_text)
            
            return ParsedDocument(
                text=full_text,
                file_path=file_path,
                file_type='pdf',
                page_count=page_count,
                structure=structure,
                metadata={'parser': 'pymupdf'}
            )
            
        except (OCRRequiredError, ParsingError):
            raise
        except Exception as e:
            # Try pdfplumber fallback for complex tables
            self.logger.warning(
                f"PyMuPDF parsing failed, trying pdfplumber fallback: {e}"
            )
            return self._parse_pdf_with_pdfplumber(file_path)
    
    def _parse_pdf_with_pdfplumber(self, file_path: str) -> ParsedDocument:
        """
        Parse PDF using pdfplumber (fallback for complex tables).
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            ParsedDocument with extracted text and structure
            
        Raises:
            OCRRequiredError: If PDF is scanned without text layer
            ParsingError: If extraction fails
        """
        try:
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                
                # Enforce page limit
                if page_count > self.MAX_PAGES:
                    raise ParsingError(
                        f"Document exceeds maximum page limit of {self.MAX_PAGES} pages "
                        f"(found {page_count} pages)"
                    )
                
                # Extract text from all pages
                text_parts = []
                has_text = False
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    
                    if page_text and page_text.strip():
                        has_text = True
                        text_parts.append(page_text)
                
                # Check if PDF has text layer
                if not has_text:
                    raise OCRRequiredError(
                        "OCR is not supported; please provide a text-based PDF"
                    )
                
                full_text = '\n'.join(text_parts)
                
                # Extract structure
                structure = self._extract_pdf_structure(file_path, full_text)
                
                return ParsedDocument(
                    text=full_text,
                    file_path=file_path,
                    file_type='pdf',
                    page_count=page_count,
                    structure=structure,
                    metadata={'parser': 'pdfplumber'}
                )
                
        except (OCRRequiredError, ParsingError):
            raise
        except Exception as e:
            raise ParsingError(f"Failed to parse PDF with pdfplumber: {e}")
    
    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """
        Parse DOCX document using python-docx.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            ParsedDocument with extracted text and structure
            
        Raises:
            ParsingError: If extraction fails
        """
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            full_text = '\n'.join(text_parts)
            
            # Extract structure
            structure = self._extract_docx_structure(doc)
            
            # Estimate page count (rough approximation)
            page_count = max(1, len(full_text) // 3000)
            
            # Enforce page limit
            if page_count > self.MAX_PAGES:
                raise ParsingError(
                    f"Document exceeds maximum page limit of {self.MAX_PAGES} pages "
                    f"(estimated {page_count} pages)"
                )
            
            return ParsedDocument(
                text=full_text,
                file_path=file_path,
                file_type='docx',
                page_count=page_count,
                structure=structure,
                metadata={'parser': 'python-docx'}
            )
            
        except ParsingError:
            raise
        except Exception as e:
            raise ParsingError(f"Failed to parse DOCX: {e}")
    
    def _parse_txt(self, file_path: str) -> ParsedDocument:
        """
        Parse TXT file with encoding detection.
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            ParsedDocument with extracted text and structure
            
        Raises:
            ParsingError: If extraction fails
        """
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
            
            # Read file with detected encoding
            with open(file_path, 'r', encoding=encoding) as f:
                full_text = f.read()
            
            # Extract structure
            structure = self._extract_txt_structure(full_text)
            
            # Estimate page count
            page_count = max(1, len(full_text) // 3000)
            
            # Enforce page limit
            if page_count > self.MAX_PAGES:
                raise ParsingError(
                    f"Document exceeds maximum page limit of {self.MAX_PAGES} pages "
                    f"(estimated {page_count} pages)"
                )
            
            return ParsedDocument(
                text=full_text,
                file_path=file_path,
                file_type='txt',
                page_count=page_count,
                structure=structure,
                metadata={'parser': 'txt', 'encoding': encoding}
            )
            
        except ParsingError:
            raise
        except Exception as e:
            raise ParsingError(f"Failed to parse TXT: {e}")
    
    def _extract_pdf_structure(self, file_path: str, text: str) -> DocumentStructure:
        """
        Extract structure from PDF text.
        
        Args:
            file_path: Path to PDF file
            text: Extracted text
            
        Returns:
            DocumentStructure with headings, sections, paragraphs
        """
        # Simple heuristic-based structure extraction
        lines = text.split('\n')
        headings = []
        sections = []
        paragraphs = []
        
        current_pos = 0
        current_section_title = None
        current_section_start = 0
        current_section_content = []
        
        for line in lines:
            line_stripped = line.strip()
            line_len = len(line) + 1  # +1 for newline
            
            # Detect headings (markdown style, all caps, numbered, or ending with colon)
            if line_stripped and len(line_stripped) < 100:
                is_heading = (
                    line_stripped.startswith('#') or  # Markdown heading
                    line_stripped.isupper() or
                    (line_stripped[0].isdigit() and '.' in line_stripped[:5]) or
                    line_stripped.endswith(':')
                )
                
                if is_heading:
                    # Save previous section if exists
                    if current_section_title:
                        section_content = '\n'.join(current_section_content)
                        sections.append(Section(
                            title=current_section_title,
                            content=section_content,
                            start_pos=current_section_start,
                            end_pos=current_pos,
                            subsections=[]
                        ))
                    
                    # Determine heading level
                    heading_level = 1
                    heading_text = line_stripped
                    if line_stripped.startswith('#'):
                        # Count leading # characters for markdown heading level
                        heading_level = len(line_stripped) - len(line_stripped.lstrip('#'))
                        heading_text = line_stripped.lstrip('#').strip()
                    
                    # Start new section
                    heading = Heading(
                        level=heading_level,
                        text=heading_text,
                        start_pos=current_pos,
                        end_pos=current_pos + line_len
                    )
                    headings.append(heading)
                    
                    current_section_title = heading_text
                    current_section_start = current_pos
                    current_section_content = []
                else:
                    current_section_content.append(line)
            else:
                current_section_content.append(line)
            
            # Track paragraphs
            if line_stripped:
                para_start = current_pos
                para_end = current_pos + line_len
                paragraphs.append(Paragraph(
                    text=line_stripped,
                    start_pos=para_start,
                    end_pos=para_end
                ))
            
            current_pos += line_len
        
        # Save last section - ensure end_pos doesn't exceed text length
        if current_section_title:
            section_content = '\n'.join(current_section_content)
            sections.append(Section(
                title=current_section_title,
                content=section_content,
                start_pos=current_section_start,
                end_pos=min(current_pos, len(text)),
                subsections=[]
            ))
        
        return DocumentStructure(
            headings=headings,
            sections=sections,
            paragraphs=paragraphs
        )
    
    def _extract_docx_structure(self, doc: DocxDocument) -> DocumentStructure:
        """
        Extract structure from DOCX document.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            DocumentStructure with headings, sections, paragraphs
        """
        headings = []
        sections = []
        paragraphs = []
        
        # First pass: build the actual text to get accurate positions
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        full_text = '\n'.join(text_parts)
        
        current_pos = 0
        current_section_title = None
        current_section_start = 0
        current_section_content = []
        
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if not para_text:
                continue
                
            para_len = len(para.text) + 1  # +1 for newline
            
            # Check if paragraph is a heading
            if para.style.name.startswith('Heading'):
                # Extract heading level
                try:
                    level = int(para.style.name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                
                # Save previous section if exists
                if current_section_title:
                    section_content = '\n'.join(current_section_content)
                    sections.append(Section(
                        title=current_section_title,
                        content=section_content,
                        start_pos=current_section_start,
                        end_pos=current_pos,
                        subsections=[]
                    ))
                
                # Add heading
                heading = Heading(
                    level=level,
                    text=para_text,
                    start_pos=current_pos,
                    end_pos=current_pos + len(para.text)
                )
                headings.append(heading)
                
                # Start new section
                current_section_title = para_text
                current_section_start = current_pos
                current_section_content = []
            else:
                current_section_content.append(para.text)
            
            # Track paragraphs
            paragraphs.append(Paragraph(
                text=para_text,
                start_pos=current_pos,
                end_pos=current_pos + len(para.text)
            ))
            
            current_pos += para_len
        
        # Save last section - ensure end_pos doesn't exceed text length
        if current_section_title:
            section_content = '\n'.join(current_section_content)
            sections.append(Section(
                title=current_section_title,
                content=section_content,
                start_pos=current_section_start,
                end_pos=min(current_pos, len(full_text)),
                subsections=[]
            ))
        
        return DocumentStructure(
            headings=headings,
            sections=sections,
            paragraphs=paragraphs
        )
    
    def _extract_txt_structure(self, text: str) -> DocumentStructure:
        """
        Extract structure from plain text.
        
        Args:
            text: Plain text content
            
        Returns:
            DocumentStructure with headings, sections, paragraphs
        """
        # Use same heuristic as PDF
        return self._extract_pdf_structure('', text)
    
    def extract_structure(self, parsed_doc: ParsedDocument) -> DocumentStructure:
        """
        Extract headings, sections, and hierarchy from parsed document.
        
        Args:
            parsed_doc: Previously parsed document
            
        Returns:
            DocumentStructure with extracted hierarchy
        """
        # Structure is already extracted during parsing
        return parsed_doc.structure
