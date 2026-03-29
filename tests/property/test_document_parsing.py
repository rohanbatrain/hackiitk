"""
Property-based tests for document parsing.

Tests universal properties that should hold for all valid documents
across PDF, DOCX, and TXT formats.
"""

import tempfile
from pathlib import Path
import pytest
from hypothesis import given, strategies as st, settings, assume
import fitz  # PyMuPDF
from docx import Document as DocxDocument

from ingestion.document_parser import (
    DocumentParser,
    UnsupportedFormatError,
    ParsingError,
)
from models.domain import ParsedDocument


# ============================================================================
# Hypothesis Strategies for Document Generation
# ============================================================================


@st.composite
def valid_text_content(draw):
    """Generate valid text content for documents."""
    # Generate 1-10 sections
    num_sections = draw(st.integers(min_value=1, max_value=10))
    sections = []
    
    for i in range(num_sections):
        heading = draw(st.text(min_size=5, max_size=50, alphabet=st.characters(
            whitelist_categories=('Lu', 'Ll', 'Nd', 'Zs'), min_codepoint=32, max_codepoint=126
        )))
        content = draw(st.text(min_size=20, max_size=200, alphabet=st.characters(
            whitelist_categories=('Lu', 'Ll', 'Nd', 'Zs', 'Po'), min_codepoint=32, max_codepoint=126
        )))
        sections.append(f"{i+1}. {heading}\n{content}\n\n")
    
    return ''.join(sections)


@st.composite
def valid_pdf_document(draw, temp_dir):
    """Generate a valid PDF document with text."""
    content = draw(valid_text_content())
    # Limit pages to avoid exceeding 100-page limit
    num_pages = draw(st.integers(min_value=1, max_value=10))
    
    pdf_path = temp_dir / f"test_{draw(st.integers(min_value=0, max_value=10000))}.pdf"
    doc = fitz.open()
    
    # Split content across pages
    lines = content.split('\n')
    lines_per_page = max(1, len(lines) // num_pages)
    
    for page_num in range(num_pages):
        page = doc.new_page()
        start_idx = page_num * lines_per_page
        end_idx = start_idx + lines_per_page
        page_lines = lines[start_idx:end_idx]
        
        y_pos = 72
        for line in page_lines:
            if line.strip():
                page.insert_text((72, y_pos), line)
                y_pos += 15
    
    doc.save(str(pdf_path))
    doc.close()
    
    return pdf_path


@st.composite
def valid_docx_document(draw, temp_dir):
    """Generate a valid DOCX document with structure."""
    content = draw(valid_text_content())
    
    docx_path = temp_dir / f"test_{draw(st.integers(min_value=0, max_value=10000))}.docx"
    doc = DocxDocument()
    
    # Parse content into sections
    sections = content.split('\n\n')
    for section in sections:
        if section.strip():
            lines = section.split('\n')
            if lines:
                # First line as heading if it looks like one
                if lines[0].strip() and (lines[0][0].isdigit() or len(lines[0]) < 50):
                    doc.add_heading(lines[0].strip(), level=2)
                    for line in lines[1:]:
                        if line.strip():
                            doc.add_paragraph(line.strip())
                else:
                    for line in lines:
                        if line.strip():
                            doc.add_paragraph(line.strip())
    
    doc.save(str(docx_path))
    
    return docx_path


@st.composite
def valid_txt_document(draw, temp_dir):
    """Generate a valid TXT document."""
    content = draw(valid_text_content())
    
    txt_path = temp_dir / f"test_{draw(st.integers(min_value=0, max_value=10000))}.txt"
    txt_path.write_text(content, encoding='utf-8')
    
    return txt_path


@st.composite
def unsupported_format_file(draw, temp_dir):
    """Generate a file with unsupported format."""
    extension = draw(st.sampled_from(['json', 'xml', 'csv', 'xlsx', 'html']))
    content = draw(st.text(min_size=10, max_size=100))
    
    file_path = temp_dir / f"test_{draw(st.integers(min_value=0, max_value=10000))}.{extension}"
    file_path.write_text(content, encoding='utf-8')
    
    return file_path, extension


# ============================================================================
# Property 3: Multi-Format Document Parsing
# **Validates: Requirements 2.1, 2.4, 2.5, 2.7**
# ============================================================================


@given(st.data())
@settings(max_examples=20, deadline=5000)
def test_property_3_multi_format_parsing_pdf(data):
    """
    Property 3: Multi-Format Document Parsing (PDF)
    
    For any valid PDF file, the DocumentParser shall successfully extract
    text content while preserving structural boundaries.
    
    **Validates: Requirements 2.1, 2.7**
    """
    parser = DocumentParser()
    
    with tempfile.TemporaryDirectory() as tmpdir:
        temp_dir = Path(tmpdir)
        pdf_path = data.draw(valid_pdf_document(temp_dir))
        
        # Parse the document
        result = parser.parse(str(pdf_path), 'pdf')
        
        # Verify ParsedDocument is returned
        assert isinstance(result, ParsedDocument)
        
        # Verify file type is correct
        assert result.file_type == 'pdf'
        
        # Verify text was extracted
        assert isinstance(result.text, str)
        assert len(result.text) > 0
        
        # Verify structure is preserved
        assert result.structure is not None
        assert hasattr(result.structure, 'headings')
        assert hasattr(result.structure, 'sections')
        assert hasattr(result.structure, 'paragraphs')
        
        # Verify page count is reasonable
        assert result.page_count > 0
        assert result.page_count <= 100


@given(st.data())
@settings(max_examples=20, deadline=5000)
def test_property_3_multi_format_parsing_docx(data):
    """
    Property 3: Multi-Format Document Parsing (DOCX)
    
    For any valid DOCX file, the DocumentParser shall successfully extract
    text content while preserving structural boundaries.
    
    **Validates: Requirements 2.4, 2.7**
    """
    parser = DocumentParser()
    
    with tempfile.TemporaryDirectory() as tmpdir:
        temp_dir = Path(tmpdir)
        docx_path = data.draw(valid_docx_document(temp_dir))
        
        # Parse the document
        result = parser.parse(str(docx_path), 'docx')
        
        # Verify ParsedDocument is returned
        assert isinstance(result, ParsedDocument)
        
        # Verify file type is correct
        assert result.file_type == 'docx'
        
        # Verify text was extracted
        assert isinstance(result.text, str)
        assert len(result.text) > 0
        
        # Verify structure is preserved
        assert result.structure is not None
        assert hasattr(result.structure, 'headings')
        assert hasattr(result.structure, 'sections')
        
        # Verify page count is reasonable
        assert result.page_count > 0
        assert result.page_count <= 100


@given(st.data())
@settings(max_examples=20, deadline=5000)
def test_property_3_multi_format_parsing_txt(data):
    """
    Property 3: Multi-Format Document Parsing (TXT)
    
    For any valid TXT file, the DocumentParser shall successfully extract
    text content while preserving structural boundaries.
    
    **Validates: Requirements 2.5, 2.7**
    """
    parser = DocumentParser()
    
    with tempfile.TemporaryDirectory() as tmpdir:
        temp_dir = Path(tmpdir)
        txt_path = data.draw(valid_txt_document(temp_dir))
        
        # Parse the document
        result = parser.parse(str(txt_path), 'txt')
        
        # Verify ParsedDocument is returned
        assert isinstance(result, ParsedDocument)
        
        # Verify file type is correct
        assert result.file_type == 'txt'
        
        # Verify text was extracted
        assert isinstance(result.text, str)
        assert len(result.text) > 0
        
        # Verify structure is preserved
        assert result.structure is not None
        assert hasattr(result.structure, 'headings')
        assert hasattr(result.structure, 'sections')
        
        # Verify page count is reasonable
        assert result.page_count > 0
        assert result.page_count <= 100


# ============================================================================
# Property 4: Unsupported Format Rejection
# **Validates: Requirements 2.6**
# ============================================================================


@given(st.data())
@settings(max_examples=20, deadline=5000)
def test_property_4_unsupported_format_rejection(data):
    """
    Property 4: Unsupported Format Rejection
    
    For any file with an unsupported format (not PDF, DOCX, or TXT),
    the DocumentParser shall return an error message specifying the
    supported formats.
    
    **Validates: Requirements 2.6**
    """
    parser = DocumentParser()
    
    with tempfile.TemporaryDirectory() as tmpdir:
        temp_dir = Path(tmpdir)
        file_path, extension = data.draw(unsupported_format_file(temp_dir))
        
        # Attempt to parse unsupported format
        with pytest.raises(UnsupportedFormatError) as exc_info:
            parser.parse(str(file_path), extension)
        
        error_msg = str(exc_info.value)
        
        # Verify error message contains format information
        assert 'Unsupported format' in error_msg or 'not supported' in error_msg
        
        # Verify error message lists supported formats
        assert 'pdf' in error_msg.lower()
        assert 'docx' in error_msg.lower()
        assert 'txt' in error_msg.lower()


# ============================================================================
# Property 5: Graceful Parsing Failure
# **Validates: Requirements 2.9**
# ============================================================================


@given(st.sampled_from(['pdf', 'docx', 'txt']))
@settings(max_examples=15, deadline=5000)
def test_property_5_graceful_parsing_failure(file_type):
    """
    Property 5: Graceful Parsing Failure
    
    For any document where text extraction fails, the DocumentParser shall
    log the specific error and terminate gracefully without crashing.
    
    **Validates: Requirements 2.9**
    """
    parser = DocumentParser()
    
    with tempfile.TemporaryDirectory() as tmpdir:
        temp_dir = Path(tmpdir)
        
        # Create a corrupted file
        corrupted_path = temp_dir / f"corrupted.{file_type}"
        corrupted_path.write_text("This is not a valid file content for the format")
        
        # Attempt to parse corrupted file
        try:
            result = parser.parse(str(corrupted_path), file_type)
            # If parsing succeeds (e.g., for TXT), that's also acceptable
            assert isinstance(result, ParsedDocument)
        except ParsingError as e:
            # Verify error is raised gracefully with a message
            assert str(e)  # Error message exists
            assert 'Failed to parse' in str(e) or 'error' in str(e).lower()
        except Exception as e:
            # Any other exception should be a ParsingError or acceptable error
            # This ensures no unhandled crashes
            assert isinstance(e, (ParsingError, UnsupportedFormatError, OSError))


# ============================================================================
# Additional Property Tests
# ============================================================================


@given(st.data())
@settings(max_examples=15, deadline=5000)
def test_property_structure_consistency(data):
    """
    Property: Structure Consistency
    
    For any parsed document, the structure positions (start_pos, end_pos)
    shall be consistent and non-overlapping.
    """
    parser = DocumentParser()
    
    with tempfile.TemporaryDirectory() as tmpdir:
        temp_dir = Path(tmpdir)
        
        # Test with different formats
        format_choice = data.draw(st.sampled_from(['pdf', 'docx', 'txt']))
        
        if format_choice == 'pdf':
            file_path = data.draw(valid_pdf_document(temp_dir))
        elif format_choice == 'docx':
            file_path = data.draw(valid_docx_document(temp_dir))
        else:
            file_path = data.draw(valid_txt_document(temp_dir))
        
        result = parser.parse(str(file_path), format_choice)
        
        # Verify structure positions are valid
        for heading in result.structure.headings:
            assert heading.start_pos >= 0
            assert heading.end_pos > heading.start_pos
            assert heading.end_pos <= len(result.text)
        
        for section in result.structure.sections:
            assert section.start_pos >= 0
            assert section.end_pos > section.start_pos
            assert section.end_pos <= len(result.text)


@given(st.data())
@settings(max_examples=15, deadline=5000)
def test_property_metadata_completeness(data):
    """
    Property: Metadata Completeness
    
    For any parsed document, the metadata shall contain parser information.
    """
    parser = DocumentParser()
    
    with tempfile.TemporaryDirectory() as tmpdir:
        temp_dir = Path(tmpdir)
        
        # Test with different formats
        format_choice = data.draw(st.sampled_from(['pdf', 'docx', 'txt']))
        
        if format_choice == 'pdf':
            file_path = data.draw(valid_pdf_document(temp_dir))
        elif format_choice == 'docx':
            file_path = data.draw(valid_docx_document(temp_dir))
        else:
            file_path = data.draw(valid_txt_document(temp_dir))
        
        result = parser.parse(str(file_path), format_choice)
        
        # Verify metadata exists and contains parser info
        assert result.metadata is not None
        assert isinstance(result.metadata, dict)
        assert 'parser' in result.metadata
        assert result.metadata['parser'] in ['pymupdf', 'pdfplumber', 'python-docx', 'txt']
