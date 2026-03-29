"""
Unit tests for Document Parser.

Tests PDF, DOCX, and TXT parsing with various scenarios including
OCR detection, format validation, and structure preservation.

**Validates: Requirements 2.1-2.9**
"""

import pytest
from pathlib import Path
import tempfile
import os

# PDF creation for testing
import fitz  # PyMuPDF
from docx import Document as DocxDocument

from ingestion.document_parser import (
    DocumentParser,
    UnsupportedFormatError,
    OCRRequiredError,
    ParsingError,
)
from models.domain import ParsedDocument


@pytest.fixture
def parser():
    """Create DocumentParser instance."""
    return DocumentParser()


@pytest.fixture
def temp_dir():
    """Create temporary directory for test files."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield Path(tmpdir)


@pytest.fixture
def sample_pdf_with_text(temp_dir):
    """Create a sample PDF with text content."""
    pdf_path = temp_dir / "sample_policy.pdf"
    doc = fitz.open()
    
    # Add pages with text
    page1 = doc.new_page()
    page1.insert_text((72, 72), "Information Security Policy\n\n")
    page1.insert_text((72, 100), "1. Purpose\n")
    page1.insert_text((72, 120), "This policy establishes the framework for protecting organizational information assets.")
    
    page2 = doc.new_page()
    page2.insert_text((72, 72), "2. Scope\n")
    page2.insert_text((72, 92), "This policy applies to all employees, contractors, and third parties.")
    
    doc.save(str(pdf_path))
    doc.close()
    
    return pdf_path


@pytest.fixture
def sample_pdf_with_table(temp_dir):
    """Create a sample PDF with complex table (for pdfplumber fallback test)."""
    pdf_path = temp_dir / "policy_with_table.pdf"
    doc = fitz.open()
    
    page = doc.new_page()
    page.insert_text((72, 72), "Risk Assessment Matrix\n\n")
    page.insert_text((72, 100), "Risk Level | Impact | Likelihood\n")
    page.insert_text((72, 120), "High       | 5      | 5\n")
    page.insert_text((72, 140), "Medium     | 3      | 3\n")
    page.insert_text((72, 160), "Low        | 1      | 1\n")
    
    doc.save(str(pdf_path))
    doc.close()
    
    return pdf_path


@pytest.fixture
def sample_pdf_empty(temp_dir):
    """Create an empty PDF (simulates OCR-required PDF)."""
    pdf_path = temp_dir / "scanned_policy.pdf"
    doc = fitz.open()
    
    # Add page without text (simulates scanned image)
    doc.new_page()
    
    doc.save(str(pdf_path))
    doc.close()
    
    return pdf_path


@pytest.fixture
def sample_docx(temp_dir):
    """Create a sample DOCX with structure."""
    docx_path = temp_dir / "sample_policy.docx"
    doc = DocxDocument()
    
    # Add heading and content
    doc.add_heading('Information Security Policy', level=1)
    doc.add_paragraph('This document establishes security requirements.')
    
    doc.add_heading('1. Purpose', level=2)
    doc.add_paragraph('The purpose is to protect organizational assets.')
    
    doc.add_heading('2. Scope', level=2)
    doc.add_paragraph('This policy applies to all employees.')
    
    doc.save(str(docx_path))
    
    return docx_path


@pytest.fixture
def sample_txt(temp_dir):
    """Create a sample TXT file."""
    txt_path = temp_dir / "sample_policy.txt"
    content = """Information Security Policy

1. Purpose
This policy establishes the framework for protecting organizational information assets.

2. Scope
This policy applies to all employees, contractors, and third parties.

3. Risk Management
The organization shall conduct regular risk assessments to identify threats.
"""
    txt_path.write_text(content, encoding='utf-8')
    
    return txt_path


@pytest.fixture
def large_pdf(temp_dir):
    """Create a PDF exceeding 100-page limit."""
    pdf_path = temp_dir / "large_policy.pdf"
    doc = fitz.open()
    
    # Add 101 pages
    for i in range(101):
        page = doc.new_page()
        page.insert_text((72, 72), f"Page {i + 1}")
    
    doc.save(str(pdf_path))
    doc.close()
    
    return pdf_path


# ============================================================================
# Test PDF Extraction
# ============================================================================


def test_parse_pdf_with_text(parser, sample_pdf_with_text):
    """Test PDF extraction with sample policy PDF."""
    result = parser.parse(str(sample_pdf_with_text), 'pdf')
    
    assert isinstance(result, ParsedDocument)
    assert result.file_type == 'pdf'
    assert result.page_count == 2
    assert 'Information Security Policy' in result.text
    assert '1. Purpose' in result.text
    assert '2. Scope' in result.text
    assert len(result.structure.headings) > 0
    assert result.metadata['parser'] == 'pymupdf'


def test_parse_pdf_auto_detect_format(parser, sample_pdf_with_text):
    """Test PDF parsing with auto-detected file type."""
    result = parser.parse(str(sample_pdf_with_text))
    
    assert result.file_type == 'pdf'
    assert result.page_count == 2


def test_parse_pdf_structure_preservation(parser, sample_pdf_with_text):
    """Test that PDF parsing preserves structural boundaries."""
    result = parser.parse(str(sample_pdf_with_text), 'pdf')
    
    # Check structure extraction
    assert result.structure is not None
    assert len(result.structure.headings) > 0
    
    # Verify headings were detected
    heading_texts = [h.text for h in result.structure.headings]
    assert any('Purpose' in text for text in heading_texts)


# ============================================================================
# Test pdfplumber Fallback
# ============================================================================


def test_pdfplumber_fallback_with_complex_table(parser, sample_pdf_with_table):
    """Test pdfplumber fallback with complex table PDF."""
    # This test verifies that the parser can handle PDFs with tables
    result = parser.parse(str(sample_pdf_with_table), 'pdf')
    
    assert isinstance(result, ParsedDocument)
    assert result.file_type == 'pdf'
    assert 'Risk Assessment Matrix' in result.text
    # Parser should be either pymupdf or pdfplumber
    assert result.metadata['parser'] in ['pymupdf', 'pdfplumber']


# ============================================================================
# Test OCR Detection
# ============================================================================


def test_ocr_required_pdf_rejection(parser, sample_pdf_empty):
    """Test OCR-required PDF rejection."""
    with pytest.raises(OCRRequiredError) as exc_info:
        parser.parse(str(sample_pdf_empty), 'pdf')
    
    assert 'OCR is not supported' in str(exc_info.value)
    assert 'text-based PDF' in str(exc_info.value)


# ============================================================================
# Test DOCX Parsing
# ============================================================================


def test_parse_docx_structure_preservation(parser, sample_docx):
    """Test DOCX structure preservation."""
    result = parser.parse(str(sample_docx), 'docx')
    
    assert isinstance(result, ParsedDocument)
    assert result.file_type == 'docx'
    assert 'Information Security Policy' in result.text
    assert '1. Purpose' in result.text
    assert '2. Scope' in result.text
    
    # Check structure extraction
    assert result.structure is not None
    assert len(result.structure.headings) > 0
    
    # Verify heading levels were preserved
    heading_levels = [h.level for h in result.structure.headings]
    assert 1 in heading_levels  # Level 1 heading exists
    assert 2 in heading_levels  # Level 2 headings exist
    
    assert result.metadata['parser'] == 'python-docx'


# ============================================================================
# Test TXT Parsing
# ============================================================================


def test_parse_txt_file_reading(parser, sample_txt):
    """Test TXT file reading."""
    result = parser.parse(str(sample_txt), 'txt')
    
    assert isinstance(result, ParsedDocument)
    assert result.file_type == 'txt'
    assert 'Information Security Policy' in result.text
    assert '1. Purpose' in result.text
    assert '2. Scope' in result.text
    assert '3. Risk Management' in result.text
    
    # Check encoding was detected
    assert 'encoding' in result.metadata
    assert result.metadata['parser'] == 'txt'


def test_parse_txt_encoding_detection(parser, temp_dir):
    """Test TXT parsing with different encoding."""
    txt_path = temp_dir / "utf16_policy.txt"
    content = "Information Security Policy\n\nThis is a test document."
    txt_path.write_text(content, encoding='utf-16')
    
    result = parser.parse(str(txt_path), 'txt')
    
    assert isinstance(result, ParsedDocument)
    assert 'Information Security Policy' in result.text


# ============================================================================
# Test Unsupported Format
# ============================================================================


def test_unsupported_format_error(parser, temp_dir):
    """Test unsupported format error."""
    # Create an unsupported file type
    unsupported_path = temp_dir / "policy.xlsx"
    unsupported_path.write_text("dummy content")
    
    with pytest.raises(UnsupportedFormatError) as exc_info:
        parser.parse(str(unsupported_path), 'xlsx')
    
    error_msg = str(exc_info.value)
    assert 'Unsupported format' in error_msg
    assert 'pdf' in error_msg.lower()
    assert 'docx' in error_msg.lower()
    assert 'txt' in error_msg.lower()


def test_unsupported_format_with_auto_detect(parser, temp_dir):
    """Test unsupported format error with auto-detection."""
    unsupported_path = temp_dir / "policy.json"
    unsupported_path.write_text('{"test": "data"}')
    
    with pytest.raises(UnsupportedFormatError):
        parser.parse(str(unsupported_path))


# ============================================================================
# Test 100-Page Limit
# ============================================================================


def test_100_page_limit_enforcement(parser, large_pdf):
    """Test 100-page limit enforcement."""
    with pytest.raises(ParsingError) as exc_info:
        parser.parse(str(large_pdf), 'pdf')
    
    error_msg = str(exc_info.value)
    assert 'exceeds maximum page limit' in error_msg
    assert '100' in error_msg


# ============================================================================
# Test Error Handling
# ============================================================================


def test_file_not_found_error(parser):
    """Test error when file doesn't exist."""
    with pytest.raises(ParsingError) as exc_info:
        parser.parse('/nonexistent/path/policy.pdf', 'pdf')
    
    assert 'File not found' in str(exc_info.value)


def test_corrupted_pdf_graceful_failure(parser, temp_dir):
    """Test graceful failure with corrupted PDF."""
    corrupted_path = temp_dir / "corrupted.pdf"
    corrupted_path.write_text("This is not a valid PDF file")
    
    with pytest.raises(ParsingError):
        parser.parse(str(corrupted_path), 'pdf')


def test_corrupted_docx_graceful_failure(parser, temp_dir):
    """Test graceful failure with corrupted DOCX."""
    corrupted_path = temp_dir / "corrupted.docx"
    corrupted_path.write_text("This is not a valid DOCX file")
    
    with pytest.raises(ParsingError):
        parser.parse(str(corrupted_path), 'docx')


# ============================================================================
# Test Structure Extraction
# ============================================================================


def test_extract_structure_method(parser, sample_pdf_with_text):
    """Test extract_structure method."""
    parsed_doc = parser.parse(str(sample_pdf_with_text), 'pdf')
    
    # extract_structure should return the already-extracted structure
    structure = parser.extract_structure(parsed_doc)
    
    assert structure is not None
    assert structure == parsed_doc.structure
    assert len(structure.headings) > 0


def test_section_extraction(parser, sample_docx):
    """Test that sections are properly extracted."""
    result = parser.parse(str(sample_docx), 'docx')
    
    assert len(result.structure.sections) > 0
    
    # Check that sections have content
    for section in result.structure.sections:
        assert section.title
        assert section.start_pos >= 0
        assert section.end_pos > section.start_pos
